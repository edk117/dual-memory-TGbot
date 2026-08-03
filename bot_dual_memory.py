"""
Telegram-бот с длинной и короткой памятью.

Функции:
  • Короткая память — in-memory история диалога (deque) для осмысленной беседы.
  • Длинная память — индексация загруженных документов (PDF/TXT/DOCX) в
    векторное хранилище ChromaDB с RAG-поиском по фрагментам.

Архитектура:
  1. При загрузке документа извлекается текст, разбивается на фрагменты,
     создаются векторные представления и сохраняются в ChromaDB. Идентификатор
     документа привязывается к пользователю в SQLite.
  2. При отправке текстового сообщения бот:
     a) извлекает историю диалога из короткой памяти;
     b) при наличии активного документа находит релевантные фрагменты
        через семантический поиск (long memory);
     c) формирует запрос к LLM, включающий системный промпт, историю
        и контекст документа;
     d) сохраняет ответ в короткую память.
"""

import asyncio
import logging
import os
import re
import sqlite3
import uuid
from collections import defaultdict, deque
from pathlib import Path

import chromadb
from aiogram import Bot, Dispatcher, F, Router
from aiogram.filters import Command, CommandStart
from aiogram.types import ErrorEvent, Message
from docx import Document as DocxDocument
from dotenv import load_dotenv
from openai import AsyncOpenAI
from pypdf import PdfReader

# ─────────────────────────────────────────────────────────────────────────────
# Конфигурация
# ─────────────────────────────────────────────────────────────────────────────

load_dotenv()

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL")

# Модель для генерации ответов (chat completions API)
RESPONSE_MODEL = os.getenv("OPENAI_RESPONSE_MODEL", "gpt-4o-mini")
# Модель для создания эмбеддингов
EMBEDDING_MODEL = os.getenv(
    "OPENAI_EMBEDDING_MODEL",
    "text-embedding-3-small",
)

# Параметры короткой памяти (количество сообщений, а не пар)
SHORT_MEMORY_LIMIT = int(os.getenv("SHORT_MEMORY_LIMIT", "40"))

# Параметры разбиения текста на фрагменты
CHUNK_SIZE = 500
CHUNK_OVERLAP = 80
TOP_K = 5  # Количество релевантных фрагментов для поиска

# Поддерживаемые форматы документов
SUPPORTED_SUFFIXES = {".pdf", ".txt", ".docx"}

# Пути к файлам и директориям
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
DOCUMENTS_DIR = DATA_DIR / "documents"
CHROMA_DIR = DATA_DIR / "chroma"
DATABASE_PATH = DATA_DIR / "bot.sqlite3"

DATA_DIR.mkdir(parents=True, exist_ok=True)
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
CHROMA_DIR.mkdir(parents=True, exist_ok=True)

# Проверка обязательных переменных окружения
if not BOT_TOKEN:
    raise RuntimeError(
        "Переменная окружения BOT_TOKEN не задана. "
        "Укажите токен Telegram-бота перед запуском."
    )
if not OPENAI_API_KEY:
    raise RuntimeError(
        "Переменная окружения OPENAI_API_KEY не задана. "
        "Укажите ключ OpenAI перед запуском."
    )

# ─────────────────────────────────────────────────────────────────────────────
# Инициализация клиентов и хранилищ
# ─────────────────────────────────────────────────────────────────────────────

# AsyncOpenAI используется для всех обращений к API — чат и эмбеддинги
openai_kwargs: dict[str, str] = {"api_key": OPENAI_API_KEY}
if OPENAI_BASE_URL:
    openai_kwargs["base_url"] = OPENAI_BASE_URL
client = AsyncOpenAI(**openai_kwargs)

# ChromaDB — векторное хранилище для длинной памяти документов
chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
collection = chroma_client.get_or_create_collection(
    name="document_memory",
    metadata={"hnsw:space": "cosine"},
)

# Короткая память: deque с ограничением длины для каждого пользователя.
# Хранит попарно сообщения user/assistant в формате OpenAI.
conversation_memory: dict[int, deque[dict]] = defaultdict(
    lambda: deque(maxlen=SHORT_MEMORY_LIMIT),
)

# SQLite — привязка пользователя к текущему загруженному документу
# (сессионное состояние длинной памяти)


def initialize_database() -> None:
    """Создаёт таблицу пользовательских сессий при первом запуске."""
    connection = sqlite3.connect(DATABASE_PATH)
    try:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS user_sessions (
                user_id INTEGER PRIMARY KEY,
                document_id TEXT,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        connection.commit()
    finally:
        connection.close()


def set_current_document(user_id: int, document_id: str) -> None:
    """Привязывает документ к пользователю или обновляет привязку."""
    connection = sqlite3.connect(DATABASE_PATH)
    try:
        connection.execute(
            """
            INSERT INTO user_sessions (user_id, document_id, updated_at)
            VALUES (?, ?, CURRENT_TIMESTAMP)
            ON CONFLICT(user_id) DO UPDATE SET
                document_id = excluded.document_id,
                updated_at = CURRENT_TIMESTAMP
            """,
            (user_id, document_id),
        )
        connection.commit()
    finally:
        connection.close()


def get_current_document(user_id: int) -> str | None:
    """Возвращает ID текущего документа пользователя или None."""
    connection = sqlite3.connect(DATABASE_PATH)
    try:
        cursor = connection.execute(
            "SELECT document_id FROM user_sessions WHERE user_id = ?",
            (user_id,),
        )
        row = cursor.fetchone()
        return row[0] if row else None
    finally:
        connection.close()


def reset_current_document(user_id: int) -> None:
    """Сбрасывает текущий документ у пользователя (устанавливает NULL)."""
    connection = sqlite3.connect(DATABASE_PATH)
    try:
        connection.execute(
            """
            INSERT INTO user_sessions (user_id, document_id, updated_at)
            VALUES (?, NULL, CURRENT_TIMESTAMP)
            ON CONFLICT(user_id) DO UPDATE SET
                document_id = NULL,
                updated_at = CURRENT_TIMESTAMP
            """,
            (user_id,),
        )
        connection.commit()
    finally:
        connection.close()


# ─────────────────────────────────────────────────────────────────────────────
# Утилиты для работы с документами
# ─────────────────────────────────────────────────────────────────────────────


def sanitize_filename(filename: str) -> str:
    """Очищает имя файла от опасных символов и транслитом."""
    filename = Path(filename).name
    filename = re.sub(r"[^a-zA-Zа-яА-ЯёЁ0-9._-]", "_", filename)
    return filename or "document"


def extract_text_from_pdf(path: Path) -> str:
    """Извлекает текст из PDF-файла по страницам."""
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def extract_text_from_docx(path: Path) -> str:
    """Извлекает текст из DOCX: абзацы и таблицы."""
    document = DocxDocument(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def load_document(path: str | Path) -> str:
    """Загружает и нормализует текст из PDF, TXT или DOCX."""
    document_path = Path(path)
    suffix = document_path.suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(document_path)
    elif suffix == ".docx":
        text = extract_text_from_docx(document_path)
    elif suffix == ".txt":
        text = document_path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(
            "Поддерживаются только файлы PDF, TXT и DOCX."
        )

    # Очистка текста от артефактов кодировки и лишних пробелов
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    if not text:
        raise ValueError("В документе не найден текст.")

    return text


def split_text_into_chunks(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    """Разбивает текст на перекрывающиеся фрагменты фиксированного размера.

    Перекрытие гарантирует, что контекст не потеряется на границах фрагментов.
    Разрыв осуществляется по ближайшему пробелу для сохранения целых слов.
    """
    if chunk_size <= 0:
        raise ValueError(
            "Размер фрагмента должен быть положительным числом."
        )
    if overlap < 0 or overlap >= chunk_size:
        raise ValueError(
            "Размер перекрытия должен быть меньше размера фрагмента."
        )

    chunks: list[str] = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = min(start + chunk_size, text_length)

        if end < text_length:
            # Ищем ближайший пробел для разрыва на слове
            boundary = text.rfind(" ", start, end)
            if boundary > start + chunk_size // 2:
                end = boundary

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_length:
            break

        # Сдвиг с перекрытием
        next_start = end - overlap
        if next_start <= start:
            next_start = start + 1
        start = next_start

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Длинная память: эмбеддинги и семантический поиск
# ─────────────────────────────────────────────────────────────────────────────


async def embed_and_index_chunks(
    chunks: list[str],
    user_id: int,
    document_id: str,
    filename: str,
) -> int:
    """Создаёт эмбеддинги для фрагментов и сохраняет их в ChromaDB.

    Использует batch-запрос к OpenAI API для эффективности.
    Возвращает количество проиндексированных фрагментов.
    """
    if not chunks:
        raise ValueError(
            "Невозможно создать эмбеддинги для пустого документа."
        )

    # Пакетное получение эмбеддингов (OpenAI API поддерживает batch)
    embedding_response = await client.embeddings.create(
        model=EMBEDDING_MODEL,
        input=chunks,
    )
    embeddings = [item.embedding for item in embedding_response.data]

    if len(embeddings) != len(chunks):
        raise RuntimeError(
            "Количество эмбеддингов не совпадает с количеством фрагментов."
        )

    # Метаданные позволяют фильтровать по пользователю и документу
    ids = [
        f"{document_id}_{index}"
        for index in range(len(chunks))
    ]
    metadatas = [
        {
            "user_id": str(user_id),
            "document_id": document_id,
            "filename": filename,
            "chunk_index": index,
        }
        for index in range(len(chunks))
    ]

    # ChromaDB.add — синхронный вызов, выносим в отдельный поток
    await asyncio.to_thread(
        collection.add,
        ids=ids,
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas,
    )

    return len(chunks)


async def retrieve_relevant_chunks(
    question: str,
    user_id: int,
    document_id: str,
) -> list[str]:
    """Находит топ-K релевантных фрагментов документа по семантическому поиску.

    Поиск ограничен текущим документом конкретного пользователя
    с помощью метаданных (where-фильтр ChromaDB).
    """
    # Эмбеддинг вопроса
    embedding_response = await client.embeddings.create(
        model=EMBEDDING_MODEL,
        input=[question],
    )
    question_embedding = embedding_response.data[0].embedding

    # Запрос к ChromaDB (синхронный — выносим в поток)
    result = await asyncio.to_thread(
        collection.query,
        query_embeddings=[question_embedding],
        n_results=TOP_K,
        where={
            "$and": [
                {"user_id": str(user_id)},
                {"document_id": document_id},
            ]
        },
        include=["documents", "metadatas", "distances"],
    )

    documents = result.get("documents") or []
    if not documents:
        return []

    return [
        document
        for document in documents[0]
        if document and document.strip()
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Короткая память + RAG: генерация ответа
# ─────────────────────────────────────────────────────────────────────────────


async def generate_answer(
    question: str,
    context: list[str],
    history: list[dict],
) -> str:
    """Формирует ответ, объединяя историю диалога (short memory) и
    релевантные фрагменты документа (long memory).

    Стратегия:
      • Системный промпт адаптируется к наличию контекста.
      • История диалога передаётся как есть (контекст беседы).
      • Фрагменты документа добавляются в текущий вопрос как контекст.
    """
    messages: list[dict] = []

    # Системный промпт адаптируется к наличию контекста
    if context:
        system_content = (
            "Ты — помощник с длинной и короткой памятью. "
            "Отвечай, опираясь на релевантные фрагменты загруженного "
            "документа и историю текущей беседы. Если в документе "
            "недостаточно информации для ответа, честно скажи об этом. "
            "Поддерживай естественный язык вопроса."
        )
    else:
        system_content = (
            "Ты — дружелюбный помощник. Поддерживай осмысленную "
            "беседу, опираясь на историю сообщений."
        )
    messages.append({"role": "system", "content": system_content})

    # История диалога (короткая память)
    messages.extend(
        msg for msg in history
        if msg.get("role") in ("user", "assistant")
    )

    # Текущий вопрос с инлайн-контекстом из документа
    if context:
        joined_context = (
            "\n\n--- Релевантные фрагменты документа ---\n\n"
            .join(context)
        )
        current_message = (
            f"Релевантные фрагменты документа:\n\n{joined_context}\n\n"
            f"Вопрос:\n{question}"
        )
    else:
        current_message = question

    messages.append({"role": "user", "content": current_message})

    response = await client.chat.completions.create(
        model=RESPONSE_MODEL,
        messages=messages,
    )

    answer = response.choices[0].message.content
    if not answer:
        raise RuntimeError("OpenAI не вернул текст ответа.")

    return answer.strip()


# ─────────────────────────────────────────────────────────────────────────────
# Хэндлеры (обработчики событий)
# ─────────────────────────────────────────────────────────────────────────────

router = Router()


@router.message(CommandStart())
async def command_start_handler(message: Message) -> None:
    """Обработчик /start — приветствие и краткая справка."""
    await message.answer(
        "Привет! Я бот с длинной и короткой памятью.\n\n"
        "Что я умею:\n"
        "• Вести осмысленную беседу, помня историю диалога.\n"
        "• Отвечать на вопросы по загруженным документам "
        "(PDF, TXT, DOCX).\n\n"
        "Команды:\n"
        "/start — запуск и справка\n"
        "/help — подробная инструкция\n"
        "/new — сбросить текущий документ\n\n"
        "Чтобы задать вопрос по документу, сначала загрузите файл."
    )


@router.message(Command("help"))
async def command_help_handler(message: Message) -> None:
    """Обработчик /help — подробная справка."""
    await message.answer(
        "Как работать с ботом:\n\n"
        "1. Отправьте файл PDF, TXT или DOCX.\n"
        "2. Дождитесь сообщения об окончании индексации.\n"
        "3. Задайте вопрос — бот найдёт релевантные фрагменты "
        "в документе и ответит с учётом контекста.\n\n"
        "Без загруженного документа бот поддерживает обычную "
        "беседу, помня историю ваших сообщений.\n\n"
        "Команда /new сбрасывает текущий документ и очищает "
        "историю диалога для чистого старта.\n"
        "История диалога хранится в памяти и сбрасывается при "
        "перезапуске бота."
    )


@router.message(Command("new"))
async def command_new_handler(message: Message) -> None:
    """Обработчик /new — сброс текущего документа и истории диалога."""
    user_id = message.from_user.id

    await asyncio.to_thread(reset_current_document, user_id)

    # Очищаем короткую память для «чистого» старта
    conversation_memory.pop(user_id, None)

    await message.answer(
        "Текущий документ сброшен. Загрузите новый PDF, TXT или DOCX, "
        "или продолжайте обычную беседу."
    )


@router.message(F.document)
async def document_handler(message: Message, bot: Bot) -> None:
    """Обработчик загрузки документа — валидация и запуск индексации."""
    filename = message.document.file_name or ""
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        await message.answer(
            "Неподдерживаемый формат файла.\n"
            "Отправьте документ в формате PDF, TXT или DOCX."
        )
        return

    await message.answer(
        "Документ получен. Извлекаю текст и создаю эмбеддинги..."
    )

    try:
        await save_and_index_document(message, bot)
    except Exception:
        logger.exception("Ошибка загрузки или индексации документа")
        await message.answer(
            "Не удалось обработать документ. Проверьте, что файл "
            "не повреждён, содержит текст и имеет формат PDF, TXT "
            "или DOCX."
        )


async def save_and_index_document(message: Message, bot: Bot) -> None:
    """Полный цикл обработки документа:
    скачивание → извлечение текста → чанкинг → эмбеддинги → индексация.
    """
    if not message.document:
        raise ValueError("Сообщение не содержит документа.")

    original_filename = message.document.file_name or "document"
    filename = sanitize_filename(original_filename)
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(
            "Поддерживаются только файлы PDF, TXT и DOCX."
        )

    user_id = message.from_user.id
    document_id = uuid.uuid4().hex

    # Сохранение файла на диск в персональную директорию пользователя
    user_directory = DOCUMENTS_DIR / str(user_id)
    user_directory.mkdir(parents=True, exist_ok=True)
    file_path = user_directory / f"{document_id}_{filename}"

    telegram_file = await bot.get_file(message.document.file_id)
    if not telegram_file.file_path:
        raise RuntimeError(
            "Telegram не вернул путь к загруженному файлу."
        )

    await bot.download_file(
        telegram_file.file_path,
        destination=file_path,
    )

    # Извлечение текста (CPU-операция — в отдельном потоке)
    text = await asyncio.to_thread(load_document, file_path)

    # Разбиение на фрагменты и создание эмбеддингов
    chunks = split_text_into_chunks(text)
    count = await embed_and_index_chunks(
        chunks,
        user_id,
        document_id,
        filename,
    )

    # Привязка документа к пользователю в SQLite
    await asyncio.to_thread(set_current_document, user_id, document_id)

    await message.answer(
        "Документ сохранён и проиндексирован.\n"
        f"Файл: {filename}\n"
        f"Фрагментов: {count}\n\n"
        "Теперь задайте вопрос по его содержимому."
    )


@router.message(F.text)
async def text_message_handler(message: Message) -> None:
    """Главный обработчик текстовых сообщений — точка интеграции
    короткой и длинной памяти.

    Логика:
      1. Получаем историю диалога из короткой памяти (deque).
      2. Проверяем наличие текущего документа (long memory, SQLite).
      3. Если документ есть — ищем релевантные фрагменты в ChromaDB.
      4. Генерируем ответ через LLM, объединяя историю и контекст.
      5. Сохраняем ответ в короткую память.
    """
    user = message.from_user
    if user is None:
        return

    user_id = user.id
    text = message.text or ""

    # Короткая память: копируем историю (без изменения оригинала deque)
    history: list[dict] = list(conversation_memory[user_id])

    # Длинная память: определяем текущий документ пользователя
    document_id = await asyncio.to_thread(get_current_document, user_id)

    # Индикатор «печатает» для лучшего UX
    await message.bot.send_chat_action(
        chat_id=message.chat.id,
        action="typing",
    )

    # Поиск релевантных фрагментов (long memory)
    context: list[str] = []
    if document_id:
        context = await retrieve_relevant_chunks(text, user_id, document_id)

    # Генерация ответа (short memory + long memory)
    try:
        answer = await generate_answer(text, context, history)
    except Exception:
        logger.exception("Ошибка генерации ответа")
        await message.answer(
            "Произошла ошибка при обращении к OpenAI API."
        )
        return

    await message.answer(answer)

    # Сохранение взаимодействия в короткую память
    conversation_memory[user_id].append({"role": "user", "content": text})
    conversation_memory[user_id].append(
        {"role": "assistant", "content": answer}
    )


@router.message()
async def unsupported_message_handler(message: Message) -> None:
    """Обработчик неподдерживаемых типов сообщений."""
    await message.answer(
        "Я поддерживаю PDF, TXT и DOCX.\n"
        "Загрузите один из этих файлов или задайте текстовый вопрос."
    )


@router.error()
async def error_handler(event: ErrorEvent) -> bool:
    """Глобальный обработчик ошибок — логирует и уведомляет пользователя."""
    logger.error(
        "Необработанная ошибка Telegram-бота",
        exc_info=(
            type(event.exception),
            event.exception,
            event.exception.__traceback__,
        ),
    )

    update = event.update
    message = getattr(update, "message", None)

    if message:
        try:
            await message.answer(
                "Произошла внутренняя ошибка. Попробуйте повторить "
                "действие позже."
            )
        except Exception:
            logger.exception(
                "Не удалось отправить сообщение об ошибке"
            )

    return True


# ─────────────────────────────────────────────────────────────────────────────
# Точка входа
# ─────────────────────────────────────────────────────────────────────────────


async def main() -> None:
    """Инициализирует базу данных и запускает long-polling."""
    initialize_database()

    bot = Bot(token=BOT_TOKEN)
    dp = Dispatcher()
    dp.include_router(router)

    try:
        await dp.start_polling(bot)
    finally:
        await bot.session.close()


if __name__ == "__main__":
    asyncio.run(main())
